# import streamlit as st
# import pandas as pd
# import joblib
# import sqlite3
# import re
# import string
# import unicodedata
# import pymupdf  # PyMuPDF for PDF processing
# from docx import Document  # python-docx for Word processing
# import os

# DB_PATH = "job_postings.db"
# MODEL_PATH = os.path.abspath("logistic_randomforest.pkl")

# def initialize_db():
#     with sqlite3.connect(DB_PATH) as conn:
#         conn.execute('''CREATE TABLE IF NOT EXISTS postings 
#                         (id INTEGER PRIMARY KEY, 
#                         job_title TEXT, 
#                         description TEXT, 
#                         prediction INTEGER)''')

# def fetch_stored_jobs():
#     with sqlite3.connect(DB_PATH) as conn:
#         df = pd.read_sql_query('SELECT job_title, description, prediction FROM postings', conn)
    
#     if not df.empty:
#         df = df.applymap(lambda x: unicodedata.normalize("NFKD", x).encode('utf-8', 'ignore').decode() if isinstance(x, str) else x)
    
#     return df

# def store_to_db(job_title, description, prediction):
#     with sqlite3.connect(DB_PATH) as conn:
#         conn.execute('INSERT INTO postings (job_title, description, prediction) VALUES (?, ?, ?)', 
#                      (job_title.encode('utf-8', 'ignore').decode(), 
#                       description.encode('utf-8', 'ignore').decode(), 
#                       prediction))

# def preprocess_text(text):
#     text = text.lower()
#     text = unicodedata.normalize("NFKD", text)
#     text = re.sub(r'[^\x00-\x7F]+', ' ', text)
#     text = text.translate(str.maketrans('', '', string.punctuation))
#     text = re.sub(r'\s+', ' ', text).strip()
#     return text

# def extract_text_from_pdf(file):
#     doc = pymupdf.open(stream=file.read(), filetype="pdf")
#     return "\n".join([page.get_text("text") for page in doc])

# def extract_text_from_docx(file):
#     doc = Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# if os.path.exists(MODEL_PATH):
#     model = joblib.load(MODEL_PATH)
# else:
#     st.error(f"Model file not found at: {MODEL_PATH}")
#     st.stop()

# def predict_job_real_or_fake(description):
#     preprocessed_text = preprocess_text(description)
#     return model.predict([preprocessed_text])[0]

# def main():
#     st.set_page_config(page_title="Job Authenticity Checker", layout="wide")
#     st.markdown(
#         """
#         <style>
#             .main-title {
#                 font-size: 32px;
#                 font-weight: bold;
#                 text-align: center;
#                 color: #FFD700;
#             }
#             .sidebar .sidebar-content {
#                 background: linear-gradient(to bottom, #1e1e2e, #282a36);
#             }
#             .stButton>button {
#                 background: #4CAF50;
#                 color: white;
#                 font-size: 18px;
#                 border-radius: 8px;
#                 padding: 10px;
#                 width: 100%;
#             }
#             .stTextInput, .stTextArea {
#                 border-radius: 8px;
#                 border: 1px solid #ccc;
#             }
#             .prediction-box {
#                 text-align: center;
#                 font-size: 20px;
#                 font-weight: bold;
#                 padding: 15px;
#                 border-radius: 8px;
#             }
#         </style>
#         """,
#         unsafe_allow_html=True,
#     )

#     st.markdown('<p class="main-title">🔍 Job Posting Authenticity Checker</p>', unsafe_allow_html=True)
#     st.sidebar.title("Navigation")
    
#     menu = ["Home", "Stored Jobs", "About"]
#     choice = st.sidebar.radio("Go to", menu)
    
#     if choice == "Home":
#         st.subheader("Upload Job Details")
#         uploaded_file = st.file_uploader("Upload job description (PDF/DOCX)", type=["pdf", "docx"], key="file_uploader")
#         job_title, description = "", ""
        
#         if uploaded_file:
#             file_type = uploaded_file.type
#             if "pdf" in file_type:
#                 file_content = extract_text_from_pdf(uploaded_file)
#             elif "word" in file_type or "msword" in file_type:
#                 file_content = extract_text_from_docx(uploaded_file)

#             file_lines = file_content.split("\n", 1)
#             if len(file_lines) >= 2:
#                 job_title, description = file_lines[0], file_lines[1]
#             else:
#                 st.warning("Invalid file format. Ensure the first line is the job title.")
#                 return
#         else:
#             job_title = st.text_input("Job Title")
#             description = st.text_area("Job Description", height=200)
        
#         if st.button("Predict", use_container_width=True):
#             if job_title and description:
#                 prediction = predict_job_real_or_fake(description)
#                 result_text = "✅ Job Posting is Real" if prediction == 0 else "❌ Likely Fake"
#                 color = "#4CAF50" if prediction == 0 else "#FF5733"
#                 st.markdown(f'<p class="prediction-box" style="background-color:{color}; color:white;">{result_text}</p>', unsafe_allow_html=True)
#                 store_to_db(job_title, description, prediction)
#             else:
#                 st.warning("Please enter a job title and description.")
    
#     elif choice == "Stored Jobs":
#         st.subheader("Stored Job Postings")
#         df = fetch_stored_jobs()
#         if not df.empty:
#             st.dataframe(df, use_container_width=True)
#         else:
#             st.warning("No job postings stored yet.")
    
#     elif choice == "About":
#         st.subheader("About This App")
#         st.info("This app helps in detecting fake job postings using machine learning models.")

# if __name__ == "__main__":
#     initialize_db()
#     main()

import streamlit as st
import pandas as pd
import joblib
import sqlite3
import re
import string
import unicodedata
import pymupdf  # PyMuPDF for PDF processing
from docx import Document  # python-docx for Word processing
import os
import numpy as np
import time
import random
import webbrowser

st.set_page_config(page_title="Job Authenticity Checker", layout="wide")

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_PATH = os.path.join(BASE_DIR, "logistic_newdataset (2).pkl")
VECTORIZER_PATH = os.path.join(BASE_DIR, "tfidf_vectorizer_newdataset.pkl")
DB_PATH = os.path.join(BASE_DIR, "job_postings.db")

# Initialize database
def initialize_db():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('''CREATE TABLE IF NOT EXISTS postings 
                        (id INTEGER PRIMARY KEY, 
                        job_title TEXT, 
                        description TEXT, 
                        prediction INTEGER)''')

# Fetch stored jobs
def fetch_stored_jobs():
    with sqlite3.connect(DB_PATH) as conn:
        df = pd.read_sql_query('SELECT job_title, description, prediction FROM postings', conn)
    if not df.empty:
        df = df.applymap(lambda x: unicodedata.normalize("NFKD", x).encode('utf-8', 'ignore').decode() if isinstance(x, str) else x)
    return df

# Store job prediction
def store_to_db(job_title, description, prediction):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('INSERT INTO postings (job_title, description, prediction) VALUES (?, ?, ?)', 
                     (job_title.encode('utf-8', 'ignore').decode(), 
                      description.encode('utf-8', 'ignore').decode(), 
                      prediction))

# Preprocessing function
def preprocess_text(text):
    text = text.lower()
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# Scam keyword detection function
SCAM_KEYWORDS = [
    "work from home", "no experience", "click here", "signup bonus", "quick money", "limited time",
    "urgent requirement", "get paid", "make money", "easy income", "bitcoin", "investment opportunity",
    "sms sending job", "form filling", "weekly payout", "earn instantly", "part time work without investment",
    "zero investment", "job guarantee", "work today get paid today", "data entry scam", "easy registration",
    "just 2 hours a day"
]

def contains_scam_keywords(text):
    text = text.lower()
    # Check for any keyword presence using regex for robustness
    for keyword in SCAM_KEYWORDS:
        # Use word boundaries to ensure exact matches
        if re.search(r'\b' + re.escape(keyword) + r'\b', text):
            return True
    return False

# File handlers
def extract_text_from_pdf(file):
    doc = pymupdf.open(stream=file.read(), filetype="pdf")
    return "\n".join([page.get_text("text") for page in doc])

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# Cached model and vectorizer loading
@st.cache_resource
def load_model_and_vectorizer():
    if os.path.exists(MODEL_PATH) and os.path.exists(VECTORIZER_PATH):
        model = joblib.load(MODEL_PATH)
        vectorizer = joblib.load(VECTORIZER_PATH)
        return model, vectorizer
    else:
        st.error("Model or vectorizer file not found!")
        st.stop()

model, tfidf_vectorizer = load_model_and_vectorizer()

# Predict function
def predict_job_real_or_fake(description):
    preprocessed_text = preprocess_text(description)
    transformed_text = tfidf_vectorizer.transform([preprocessed_text])
    return model.predict(transformed_text)[0]

# Search web for company links
def get_company_links(company_name):
    search_url = f"https://www.google.com/search?q={company_name.replace(' ', '+')}+company+official"
    return search_url

# Streamlit UI
def main():
    # Initialize session state
    if "show_loader" not in st.session_state:
        st.session_state.show_loader = False
    if "prediction_done" not in st.session_state:
        st.session_state.prediction_done = False
    if "job_title" not in st.session_state:
        st.session_state.job_title = ""
    if "description" not in st.session_state:
        st.session_state.description = ""
    if "result_text" not in st.session_state:
        st.session_state.result_text = ""
    if "prediction" not in st.session_state:
        st.session_state.prediction = None

    # CSS styling
    st.markdown("""<style>
        .stApp {
            background: linear-gradient(135deg, #FFE1E1 0%, #E1F5FF 100%);
            font-family: 'Poppins', sans-serif;
            min-height: 100vh;
            display: flex;
            flex-direction: column;
            align-items: center;
            padding: 20px;
        }
        .main-title {
            font-size: 100px;
            font-weight: 700;
            text-align: center;
            text-transform: uppercase;
            letter-spacing: 5px;
            display: inline-block;
            background: linear-gradient(90deg, #FF6F61 25%, #FFCA28 50%, #FF6F61 75%);
            background-size: 200% 100%;
            animation: shine 1s ease-in-out infinite;
            -webkit-background-clip: text;
            color: transparent;
        }
        @keyframes shine {
            0% { background-position: -100%; }
            100% { background-position: 100%; }
        }
        .prediction-box {
            background-color: #FFFFFF;
            color: #333;
            text-align: center;
            font-size: 22px;
            font-weight: 600;
            padding: 20px;
            border-radius: 12px;
            box-shadow: 0 6px 12px rgba(0, 0, 0, 0.1);
            margin-top: 20px;
            transition: transform 0.3s ease;
            animation: slideIn 1s ease-out;
        }
        @keyframes slideIn {
            0% { transform: translateY(30px); opacity: 0; }
            100% { transform: translateY(0); opacity: 1; }
        }
        .fullscreen-loader {
            position: fixed !important;
            top: 0 !important;
            left: 0 !important;
            width: 100vw !important;
            height: 100vh !important;
            background: rgba(255, 255, 255, 0.9) !important;
            z-index: 9999 !important;
            display: flex !important;
            justify-content: center !important;
            align-items: center !important;
        }
        .loader {
            width: 80px !important;
            aspect-ratio: 1 !important;
            border: 15px solid #ddd !important;
            border-radius: 50% !important;
            position: relative !important;
            transform: rotate(45deg) !important;
        }
        .loader::before {
            content: "" !important;
            position: absolute !important;
            inset: -15px !important;
            border-radius: 50% !important;
            border: 15px solid #514b82 !important;
            animation: l18 2s infinite linear !important;
        }
        @keyframes l18 {
            0% {clip-path:polygon(50% 50%,0 0,0 0,0 0,0 0,0 0)}
            25% {clip-path:polygon(50% 50%,0 0,100% 0,100% 0,100% 0,100% 0)}
            50% {clip-path:polygon(50% 50%,0 0,100% 0,100% 100%,100% 100%,100% 100%)}
            75% {clip-path:polygon(50% 50%,0 0,100% 0,100% 100%,0 100%,0 100%)}
            100% {clip-path:polygon(50% 50%,0 0,100% 0,100% 100%,0 100%,0 0)}
        }
    </style>""", unsafe_allow_html=True)

    # Create a placeholder for the loader
    loader_placeholder = st.empty()

    # Show loader in the placeholder if active
    if st.session_state.show_loader:
        with loader_placeholder.container():
            st.markdown('<div class="fullscreen-loader"><div class="loader"></div></div>', unsafe_allow_html=True)
            st.markdown('<p style="text-align: center; font-size: 20px; color: #514b82;">Analyzing job description...</p>', unsafe_allow_html=True)
            st.write("DEBUG: Loader is active")

    st.markdown('<p class="main-title">Job Posting Authenticity Checker</p>', unsafe_allow_html=True)

    st.sidebar.title("Navigation")
    menu = ["Home", "Stored Jobs", "About"]
    choice = st.sidebar.radio("Go to", menu)

    if choice == "Home":
        st.subheader("Upload Job Details")

        # Input fields
        uploaded_file = st.file_uploader("Upload job description (PDF/DOCX)", type=["pdf", "docx"])
        job_title, description = "", ""

        if uploaded_file:
            file_type = uploaded_file.type
            if "pdf" in file_type:
                file_content = extract_text_from_pdf(uploaded_file)
            elif "word" in file_type or "msword" in file_type:
                file_content = extract_text_from_docx(uploaded_file)

            file_lines = file_content.split("\n", 1)
            if len(file_lines) >= 2:
                job_title, description = file_lines[0], file_lines[1]
            else:
                st.warning("Invalid file format. Ensure the first line is the job title.")
                return
        else:
            job_title = st.text_input("Job Title", value=st.session_state.job_title)
            description = st.text_area("Job Description", value=st.session_state.description, height=200)

        if st.button("Predict", use_container_width=True):
            if job_title and description:
                if len(description.split()) < 30:
                    st.warning("Please enter a job description of at least 30 words.")
                else:
                    st.session_state.job_title = job_title
                    st.session_state.description = description
                    st.session_state.show_loader = True
                    st.session_state.prediction_done = False
                    st.session_state.result_text = ""
                    st.session_state.prediction = None
                    # Process prediction
                    with loader_placeholder.container():
                        st.markdown('<div class="fullscreen-loader"><div class="loader"></div></div>', unsafe_allow_html=True)
                        st.markdown('<p style="text-align: center; font-size: 20px; color: #514b82;">Analyzing job description...</p>', unsafe_allow_html=True)
                        time.sleep(random.randint(6, 8))  # Simulate processing time
                    # Check for scam keywords first
                    if contains_scam_keywords(description):
                        prediction = 1
                        result_text = "This job posting is likely fraudulent."
                        color = "#FF5733"
                        debug_message = "DEBUG: Prediction based on scam keywords detected."
                    else:
                        prediction = predict_job_real_or_fake(description)
                        result_text = "This job posting appears to be authentic." if prediction == 0 else "This job posting is likely fraudulent."
                        color = "#4CAF50" if prediction == 0 else "#FF5733"
                        debug_message = "DEBUG: Prediction based on pickle model."

                    st.session_state.result_text = result_text
                    st.session_state.prediction = prediction
                    st.session_state.show_loader = False
                    st.session_state.prediction_done = True
                    store_to_db(job_title, description, prediction)
                    loader_placeholder.empty()  # Clear the loader
                    st.write(debug_message)  # Show which prediction method was used
                    st.rerun()  # Force rerun to update UI
            else:
                st.warning("Please enter a job title and description.")

        # Display prediction if done
        if st.session_state.prediction_done:
            st.markdown(
                f'<p class="prediction-box" style="background-color:{("#4CAF50" if st.session_state.prediction == 0 else "#FF5733")}; color:white;">{st.session_state.result_text}</p>',
                unsafe_allow_html=True
            )

            if not contains_scam_keywords(st.session_state.description):
                stack_trace = [
                    "Input received and preprocessed.",
                    "Text transformed using TF-IDF Vectorizer.",
                    "Transformed input passed to Logistic Regression model.",
                    f"Model prediction: {st.session_state.prediction} → Classified as {'REAL' if st.session_state.prediction == 0 else 'FAKE'} job."
                ]
                explanation_html = "<div class='prediction-box' style='background-color:#ffffff; color:#333333; font-size:18px; margin-top:10px;'>"
                explanation_html += "<b>Stack Trace Explanation:</b><ul style='margin-top: 10px;'>"
                for step in stack_trace:
                    explanation_html += f"<li>{step}</li>"
                explanation_html += "</ul></div>"
                st.markdown(explanation_html, unsafe_allow_html=True)

            st.sidebar.subheader("🔗 Related Company Info")
            company_search_url = get_company_links(st.session_state.job_title)
            st.sidebar.markdown(f"[Search about {st.session_state.job_title}]({company_search_url})")

    elif choice == "Stored Jobs":
        st.subheader("Stored Job Postings")
        df = fetch_stored_jobs()
        if not df.empty:
            st.dataframe(df, use_container_width=True)
        else:
            st.warning("No job postings have been stored yet.")

    elif choice == "About":
        st.subheader("About This App")
        st.info("This application assists in detecting fraudulent job postings using a combination of machine learning predictions and keyword-based analysis.")

if __name__ == "__main__":
    initialize_db()
    main()

# import re
# import string
# import unicodedata
# import pymupdf  # PyMuPDF for PDF processing
# from docx import Document  # python-docx for Word processing

# # st.set_page_config(page_title="Job Authenticity Checker", layout="wide")

# # Set base directimport os
# import joblib
# import streamlit as st
# import pandas as pd
# import sqlite3ory for model and vectorizer
# BASE_DIR = r"C:\xampp_inuse\htdocs\web\formvalidation\prediction"
# MODEL_PATH = os.path.join(BASE_DIR, "logistic_newdataset (2).pkl")
# VECTORIZER_PATH = os.path.join(BASE_DIR, "tfidf_vectorizer_newdataset.pkl")
# DB_PATH = os.path.join(BASE_DIR, "job_postings.db")

# # Initialize database
# def initialize_db():
#     with sqlite3.connect(DB_PATH) as conn:
#         conn.execute('''CREATE TABLE IF NOT EXISTS postings 
#                         (id INTEGER PRIMARY KEY, 
#                         job_title TEXT, 
#                         description TEXT, 
#                         prediction INTEGER)''')

# # Load model and vectorizer
# if os.path.exists(MODEL_PATH) and os.path.exists(VECTORIZER_PATH):
#     model = joblib.load(MODEL_PATH)
#     tfidf_vectorizer = joblib.load(VECTORIZER_PATH)
# else:
#     st.error(f"🚨 Model or vectorizer file not found!\nExpected at:\n {MODEL_PATH}\n {VECTORIZER_PATH}")
#     st.stop()

# # ✅ Function to preprocess job descriptions
# def preprocess_text(text):
#     text = text.lower()
#     text = unicodedata.normalize("NFKD", text)
#     text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
#     text = text.translate(str.maketrans('', '', string.punctuation))  # Remove punctuation
#     text = re.sub(r'\s+', ' ', text).strip()  # Remove extra spaces
#     return text

# # ✅ Function to extract text from uploaded PDF/DOCX
# def extract_text_from_file(uploaded_file):
#     text = ""
#     if uploaded_file is not None:
#         file_extension = uploaded_file.name.split(".")[-1]
#         if file_extension == "pdf":
#             pdf_doc = pymupdf.open(stream=uploaded_file.read(), filetype="pdf")
#             text = " ".join([page.get_text() for page in pdf_doc])
#         elif file_extension == "docx":
#             doc = Document(uploaded_file)
#             text = " ".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def main():
#     st.markdown('<p class="main-title">🔍 Job Posting Authenticity Checker</p>', unsafe_allow_html=True)
    
#     menu = ["Home", "Stored Jobs", "About"]
#     choice = st.sidebar.radio("Go to", menu)
    
#     if choice == "Home":
#         st.subheader("Upload or Enter Job Details")
        
#         uploaded_file = st.file_uploader("Upload job description (PDF/DOCX)", type=["pdf", "docx"])
#         job_title = st.text_input("Job Title (Optional)")
#         manual_description = st.text_area("Or Enter Job Description Manually", height=200)

#         # ✅ Extract text from uploaded file if provided
#         extracted_text = extract_text_from_file(uploaded_file) if uploaded_file else ""

#         # ✅ Prioritize uploaded text, otherwise use manual input
#         final_description = extracted_text if extracted_text else manual_description

#         if st.button("Predict", use_container_width=True):
#             if final_description:  # ✅ Only require job description, job title is optional
#                 processed_text = preprocess_text(final_description)
#                 prediction = model.predict(tfidf_vectorizer.transform([processed_text]))[0]
#                 result_text = "✅ Job Posting is Real" if prediction == 0 else "❌ Likely Fake"
#                 st.markdown(f'<p style="background-color:#4CAF50; padding:10px; border-radius:8px;">{result_text}</p>', unsafe_allow_html=True)
#             else:
#                 st.warning("⚠️ Please upload a file OR enter a job description.")

# if __name__ == "__main__":
#     initialize_db()
#     main()




# /////////////////////////////////////////////////////////////////////////
# import streamlit as st
# import time
# import pandas as pd
# import random
# import joblib
# import os
# import sqlite3
# import unicodedata
# import re
# import string
# import docx
# import fitz  # PyMuPDF

# # --- Load Model & Vectorizer ---
# import joblib

# DB_PATH = "job_postings.db"
# MODEL_PATH = "formvalidation/prediction/logistic_randomforest.pkl"
# VECTORIZER_PATH = "formvalidation/prediction/tfidf_vectorizer.pkl"

# # Load the model
# try:
#     model = joblib.load(MODEL_PATH)
#     print("Model loaded successfully!")
# except Exception as e:
#     print("Error loading model:", e)

# # Load the TF-IDF vectorizer
# try:
#     tfidf_vectorizer = joblib.load(VECTORIZER_PATH)
#     print("Vectorizer loaded successfully!")
# except Exception as e:
#     print("Error loading vectorizer:", e)


# # --- Styling ---
# st.set_page_config(page_title="Job Authenticity Checker", layout="wide")

# st.markdown("""
#     <style>
#     body { background-color: #F9F9F9; font-family: 'Segoe UI', sans-serif; }
#     .main-title { font-size: 3em; font-weight: bold; color: #4A4A4A; margin-bottom: 10px; text-align: center; }
#     .subtitle { font-size: 1.2em; color: #6C757D; text-align: center; margin-bottom: 30px; }
#     .upload-box { background-color: #FFFFFF; padding: 20px; border-radius: 12px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); margin-bottom: 20px; }
#     .analyze-button button { background-color: #4CAF50; color: white; font-weight: bold; border-radius: 10px; padding: 0.75em 2em; }
#     .result-container { background-color: #FFFFFF; padding: 30px; border-radius: 15px; box-shadow: 0 8px 24px rgba(0,0,0,0.15); margin-top: 30px; }
#     .highlight-file { font-size: 0.95em; color: #555; margin-top: 10px; font-style: italic; }
#     .loading-box { display: flex; flex-direction: column; align-items: center; justify-content: center; padding: 30px; background-color: #f0f8ff; border-radius: 15px; margin-top: 20px; }
#     .fade-text { margin-top: 15px; font-size: 1.2em; color: #007BFF; animation: fadeIn 2s infinite alternate; }
#     .wave-container { display: flex; justify-content: center; gap: 8px; margin-top: 15px; }
#     .wave-bar { width: 10px; height: 20px; background: #4CAF50; border-radius: 5px; animation: wave 1.2s infinite ease-in-out; }
#     .wave-bar:nth-child(2) { animation-delay: -1.1s; }
#     .wave-bar:nth-child(3) { animation-delay: -1.0s; }
#     .wave-bar:nth-child(4) { animation-delay: -0.9s; }
#     .wave-bar:nth-child(5) { animation-delay: -0.8s; }
#     .ball { width: 12px; height: 12px; background-color: #007BFF; border-radius: 50%; margin-top: 10px; animation: bounce 1s infinite; }
#     .ball:nth-child(2) { animation-delay: 0.2s; }
#     .ball:nth-child(3) { animation-delay: 0.4s; }

#     @keyframes wave {
#         0%, 100% { transform: scaleY(1); }
#         50% { transform: scaleY(2); }
#     }
#     @keyframes bounce {
#         0%, 100% { transform: translateY(0); }
#         50% { transform: translateY(-10px); }
#     }
#     @keyframes fadeIn {
#         0% { opacity: 0.4; }
#         100% { opacity: 1; }
#     }
#     </style>
# """, unsafe_allow_html=True)

# st.markdown('<div class="main-title">🕵️‍♂️ Job Authenticity Checker</div>', unsafe_allow_html=True)
# st.markdown('<div class="subtitle">Paste a job description or upload a file.</div>', unsafe_allow_html=True)

# # --- Input Fields ---
# st.markdown('<div class="upload-box">', unsafe_allow_html=True)
# job_title = st.text_input("Job Title")
# job_description = st.text_area("Paste Job Description Here", height=200)
# uploaded_file = st.file_uploader("Or Upload a Job Description (PDF or DOCX)", type=["pdf", "docx"])
# st.markdown('</div>', unsafe_allow_html=True)

# # --- Scam Keywords Check ---
# scam_keywords = [
#     "quick money", "work from home", "no experience needed", "click here", "earn cash",
#     "urgent requirement", "100% free", "send money", "training fee", "instant hiring",
#     "no interview", "payment before joining", "guaranteed income", "easy income",
#     "limited seats", "signup bonus", "sms job", "whatsapp job", "daily payout", "money transfer"
# ]

# def contains_scam_keywords(text):
#     text = text.lower()
#     return any(keyword in text for keyword in scam_keywords)

# # --- Preprocess Function ---
# def preprocess_text(text):
#     text = unicodedata.normalize('NFKD', text)
#     text = text.lower()
#     text = re.sub(r'http\S+', '', text)
#     text = re.sub(r'<.*?>', '', text)
#     text = text.translate(str.maketrans('', '', string.punctuation))
#     text = re.sub(r'\d+', '', text)
#     text = ' '.join(text.split())
#     return text

# # --- File Text Extraction ---
# def extract_text_from_file(uploaded_file):
#     if uploaded_file is not None:
#         extension = os.path.splitext(uploaded_file.name)[1].lower()
#         if extension == ".pdf":
#             doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
#             return " ".join([page.get_text() for page in doc])
#         elif extension == ".docx":
#             doc = docx.Document(uploaded_file)
#             return " ".join([para.text for para in doc.paragraphs])
#     return ""

# # --- Create SQLite Table If Not Exists ---
# def create_db():
#     with sqlite3.connect(DB_PATH) as conn:
#         conn.execute('''CREATE TABLE IF NOT EXISTS postings (
#                             id INTEGER PRIMARY KEY AUTOINCREMENT,
#                             job_title TEXT,
#                             description TEXT,
#                             prediction INTEGER
#                         )''')

# create_db()

# # --- Analysis Button ---
# if st.button("Analyze Now"):
#     extracted_text = extract_text_from_file(uploaded_file) if uploaded_file else ""
#     final_description = extracted_text if extracted_text else job_description.strip()

#     if final_description:
#         processed_text = preprocess_text(final_description)

#         # If scam keywords detected
#         if contains_scam_keywords(processed_text):
#             prediction = 1  # Force fake
#             st.warning("⚠️ This job contains suspicious phrases and is classified as **FAKE**.")
#         else:
#             # Loading animation
#             loading_placeholder = st.empty()
#             loading_placeholder.markdown(
#                 '<div class="loading-box"><div class="wave-container"><div class="wave-bar"></div><div class="wave-bar"></div><div class="wave-bar"></div><div class="wave-bar"></div><div class="wave-bar"></div></div><span class="fade-text">Processing Your Job Check! 🎉 Stay Tuned... 🚀</span><div class="ball"></div><div class="ball"></div><div class="ball"></div></div>',
#                 unsafe_allow_html=True
#             )

#             start_time = time.time()
#             vectorized_text = tfidf_vectorizer.transform([processed_text])
#             prediction = model.predict(vectorized_text)[0]
#             elapsed_time = time.time() - start_time
#             time.sleep(max(0, 8 - elapsed_time))

#             loading_placeholder.empty()

#         # Save result to DB
#         with sqlite3.connect(DB_PATH) as conn:
#             conn.execute("INSERT INTO postings (job_title, description, prediction) VALUES (?, ?, ?)",
#                          (job_title, final_description, int(prediction)))
#             conn.commit()

#         # Display result
#         st.markdown('<div class="result-container">', unsafe_allow_html=True)
#         if prediction == 0:
#             st.success("✅ The job posting seems legitimate!")
#             st.markdown("**Next Steps:** Proceed with confidence, but always double-check details.")
#         else:
#             st.error("❌ This job posting may be fake! 🚨")
#             st.markdown("**Recommendation:** Verify the source and contact the employer directly.")
#         st.markdown('</div>', unsafe_allow_html=True)

#         if uploaded_file:
#             st.markdown(f'<div class="highlight-file">📄 {uploaded_file.name}</div>', unsafe_allow_html=True)
#     else:
#         st.warning("⚠️ Please upload a file or enter a job description.")



